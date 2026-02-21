from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align
from src.utils.helpers import style_as

def add_education(doc, school, grad_date=None, degree=None, coursework=None):
    table=doc.add_table(rows=1,cols=2)
    table.autofit=False                
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.5)

    left=table.cell(0, 0).paragraphs[0] 
    right=table.cell(0, 1).paragraphs[0] 
    school=left.add_run(school)
    style_as(school, size=12, bold=True)
    if grad_date:
        right.alignment=align.RIGHT
        gdate=right.add_run(grad_date)
        style_as(gdate, size=12, italic=True)
    DC_para=doc.add_paragraph()
    deg=DC_para.add_run(degree)
    style_as(deg, size=12, bold=False)
    crse=DC_para.add_run(f"\nRelevant coursework: {coursework}")
    style_as(crse, size=12, italic=False)
