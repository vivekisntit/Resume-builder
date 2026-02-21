from docx.shared import Pt
from utils.helpers import style_as

def add_skills(doc, prog_lang, frm_lib, duties):
    para_skill=doc.add_paragraph()
    pl=para_skill.add_run("Programming Languages: ")
    style_as(pl, size=12, bold=True)  
    plnag=para_skill.add_run(" | ".join(prog_lang))
    style_as(plnag, size=10, bold=False)
    para_skill.add_run("\n")  
    fl=para_skill.add_run("Frameworks & Libraries: ")
    style_as(fl, size=12, bold=True)
    flib=para_skill.add_run(" | ".join(frm_lib))
    style_as(flib, size=10, bold=False)
    para_skill.add_run("\n")
    rr=para_skill.add_run("Roles & Responsibilities: ")
    style_as(rr, size=12, bold=True)
    rresp=para_skill.add_run(" | ".join(duties))
    style_as(rresp, size=10, bold=False)
    para_skill.paragraph_format.space_after = Pt(0)
    para_skill.paragraph_format.space_before = Pt(0)
