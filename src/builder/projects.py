from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align
from utils.helpers import style_as

def add_project(doc, pr_cr_name, timeframe=None, description=None, framework=None, domain=None):
    table=doc.add_table(rows=1, cols=2)
    table.autofit=False                  
    table.columns[0].width = Inches(5.5)   
    table.columns[1].width = Inches(1.5)   
    left=table.cell(0, 0).paragraphs[0]
    right=table.cell(0, 1).paragraphs[0]
    PC_name=left.add_run(pr_cr_name)
    style_as(PC_name, size=12, bold=True)
    PC_name.font.underline=True
    if timeframe:
        right.alignment=align.RIGHT
        PC_time=right.add_run(timeframe)
        style_as(PC_time, size=11, italic=True)
    if description:
        for d in description:
            Discriotion_para=doc.add_paragraph(d, style="List Bullet")
            Discriotion_para.paragraph_format.left_indent = Inches(0.5)
            Discriotion_para.paragraph_format.space_before = Pt(0)
            Discriotion_para.paragraph_format.space_after = Pt(0)
    if framework or domain:
        fwork = ", ".join(framework) if framework else ""
        dmn = domain if domain else ""
        line = f"Framework: {fwork}"
        if dmn:
            line += f" | {dmn}"

        FDmn_para=doc.add_paragraph(line, style="List Bullet")
        FDmn_para.paragraph_format.left_indent = Inches(0.5)
        FDmn_para.paragraph_format.space_before = Pt(0)
        FDmn_para.paragraph_format.space_after = Pt(0)
