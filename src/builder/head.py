from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from src.utils.helpers import style_as, add_hlink

def add_sec_head(doc, title):

    Head_para=doc.add_paragraph()
    head=Head_para.add_run(title)
    style_as(head, size=15, bold=True)
    Head_para_pos=Head_para.paragraph_format
    Head_para_pos.space_before = Pt(6)
    Head_para_pos.space_after = Pt(0)
    Head_para_pos.line_spacing = Pt(14)
    
    Head_para_props = Head_para._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    p_borders.append(bottom)
    Head_para_props.append(p_borders)
    gap=doc.add_paragraph()
    gap.paragraph_format.space_before=Pt(0)
    gap.paragraph_format.space_after =Pt(6)
    gap.paragraph_format.line_spacing=Pt(12)
    
def add_head_info(doc, name, phone=None, email=None, github=None, linkedin=None):

    name1= doc.add_paragraph()
    name1.alignment = align.CENTER
    a = name1.add_run(name)
    style_as(a, size=24, bold=True)

    infoLINE = doc.add_paragraph()
    infoLINE.alignment = align.CENTER

    infoSLOT=[]
    if phone: 
        infoSLOT.append(phone)
    if email: 
        infoSLOT.append(email)
    if infoSLOT:
        b=infoLINE.add_run("  |  ".join(infoSLOT))
        style_as(b, size=10, bold=False)
    if github:
        if infoSLOT: 
            infoLINE.add_run("  |  ")
        add_hlink(infoLINE, f"https://github.com/{github}", "GitHub")
    if linkedin:
        if infoSLOT or github: 
            infoLINE.add_run("  |  ")
        add_hlink(infoLINE, f"https://www.linkedin.com/in/{linkedin}", "LinkedIn")
