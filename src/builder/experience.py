from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align
from src.utils.helpers import style_as

def add_experience(doc, job, level, employer, date_range=None, description=None):
    table=doc.add_table(rows=1, cols=2)
    table.autofit=False                   
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.5)

    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    job1=left.add_run(job)
    style_as(job1, size=12, bold=True)
    job1.font.underline=True  
    left.add_run(" | ")
    level1=left.add_run(level)
    style_as(level1, size=12, bold=True)

    if date_range:
        right.alignment=align.RIGHT
        dRange=right.add_run(date_range)
        style_as(dRange, size=11, italic=True)
    Employer_para=doc.add_paragraph()
    Employer_para.paragraph_format.space_before = Pt(0)
    Employer_para.paragraph_format.space_after = Pt(0)
    emp_data=Employer_para.add_run(employer)
    style_as(emp_data, size=11, italic=True)
    if description:
        for d in description:
            job_des=doc.add_paragraph(d, style="List Bullet")
            job_des.paragraph_format.left_indent = Inches(0.5)
            job_des.paragraph_format.space_before = Pt(0)
            job_des.paragraph_format.space_after = Pt(0)
