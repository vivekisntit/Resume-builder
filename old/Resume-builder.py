# import os
import sys
from openai import OpenAI, AuthenticationError
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as align
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# from docx2pdf import convert
# import json
def style_as(txt, size=13, bold=False, italic=False, name="Times New Roman"):
    txt.font.name=name
    txt.font.size=Pt(size)
    txt.bold=bold
    txt.italic=italic

def add_hlink(Input_para, url, Input_text, size=10, bold=False):
    part = Input_para.part
    reID = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hlink = OxmlElement('w:hyperlink')
    hlink.set(qn('r:id'), reID)

    txtRun = OxmlElement('w:r')
    txtEdits = OxmlElement('w:rPr')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF'); txtEdits.append(color)
    underline = OxmlElement('w:u'); underline.set(qn('w:val'), 'single'); txtEdits.append(underline)
    txtRun.append(txtEdits)
    txt = OxmlElement('w:t'); txt.text = Input_text; txtRun.append(txt)
    hlink.append(txtRun); Input_para._p.append(hlink)
    return hlink

def validate_api(api_key: str) -> bool:
    try:
        client = OpenAI(api_key=api_key)
        client.models.list()
        return True
    except AuthenticationError:
        return False

def add_sec_head(doc, title):

    Head_para=doc.add_paragraph()
    head=Head_para.add_run(title)
    style_as(head, size=15, bold=True)
    Head_para_pos=Head_para.paragraph_format
    Head_para_pos.space_before = Pt(6)
    Head_para_pos.space_after = Pt(0)
    Head_para_pos.line_spacing = Pt(14)
    
    Head_para_props = Head_para._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    p_borders.append(bottom)
    Head_para_props.append(p_borders)
    gap=doc.add_paragraph()
    gap.paragraph_format.space_before=Pt(0)
    gap.paragraph_format.space_after =Pt(6)
    gap.paragraph_format.line_spacing=Pt(12)
    
def add_head_info(doc, name, phone=None, email=None, github=None, linkedin=None):

    name1= doc.add_paragraph()
    name1.alignment = align.CENTER
    a = name1.add_run(name)
    style_as(a, size=24, bold=True)

    infoLINE = doc.add_paragraph()
    infoLINE.alignment = align.CENTER

    infoSLOT=[]
    if phone: 
        infoSLOT.append(phone)
    if email: 
        infoSLOT.append(email)
    if infoSLOT:
        b=infoLINE.add_run("  |  ".join(infoSLOT))
        style_as(b, size=10, bold=False)
    if github:
        if infoSLOT: 
            infoLINE.add_run("  |  ")
        add_hlink(infoLINE, f"https://github.com/{github}", "GitHub")
    if linkedin:
        if infoSLOT or github: 
            infoLINE.add_run("  |  ")
        add_hlink(infoLINE, f"https://www.linkedin.com/in/{linkedin}", "LinkedIn")
        
def get_input(user_inp):
    return input(f"{user_inp}\n> ").strip()
    
def get_dynamic_list(user_inp):
    dyn_items=[]
    print(f"Enter {user_inp} (type 'stop' to finish):")
    while True:
        item = input("> ").strip()
        if item.lower() == "stop":
            break
        if item:
            dyn_items.append(item)
    return dyn_items
    
def get_experiences():
    experiences=[]
    i=1
    while True:
        print(f"\n--- Experience {i} ---")
        exp={
            "job": get_input("Enter job title"),
            "level": get_input("Enter job level (Eg: Intern, Full-time)"),
            "employer": get_input("Enter employer name"),
            "date_range": get_input("Enter date range (Eg: Jan 2020 - Dec 2021)"),
            "description": get_dynamic_list("experience description bullets")
        }
        experiences.append(exp)
        more=get_input(f"Experience {i+1}? (y/n)").lower()
        if more!="y":
            break
        i+=1
    return experiences
    
def get_projects():
    projects=[]
    i=1
    while True:
        print(f"\n--- Project {i} ---")
        proj={
            "pr_cr_name": get_input("Enter project name"),
            "timeframe": get_input("Enter project timeframe"),
            "description": get_dynamic_list("project description bullets"),
            "framework": get_dynamic_list("Mention the frameworks/libraries which you used in this project"),
            "domain": get_input("Enter project domain (Eg: Web dev, App dev, ML engineer etc. [WRITE IN FULL FORM])")
        }
        projects.append(proj)
        more=get_input(f"Project {i+1}? (y/n)").lower()
        if more!="y":
            break
        i+=1
    return projects

def build_resume():
    resume_data={
        "name_and_contact": {
            "name": get_input("Enter your full name"),
            "phone": get_input("Enter your phone number"),
            "email": get_input("Enter your email"),
            "github": get_input("Enter your GitHub username"),
            "linkedin": get_input("Enter your LinkedIn handle")
        },

        "education": [
            {
                "school": get_input("Enter your school/college name"),
                "grad_date": get_input("Enter your graduation year"),
                "degree": get_input("Enter your degree"),
                "coursework": get_input("Enter your coursework (Eg: WebDev, DBMS, DSA, ML, Data Science etc. [WRITE IN FULL FORM] ")
            }
        ],

        "experience": get_experiences(),
        "projects": get_projects(),

        "skills": {
            "prog_lang": get_dynamic_list("Mention the programming languages you know"),
            "frm_lib": get_dynamic_list("Mention any frameworks/libraries/tools you know"),
            "duties": get_dynamic_list("Mention any duties/soft skills you have")
        }
    }
    return resume_data

def enhance_resume_data(input_data, client):
    # prompt for project enhancement
    openai_guider_0=(
        "You are a resume assistant."
        "Rewrite resume bullet points to be clear, precise and professional."
        "Keep technical terms, metrics, and achievements intact."
        "Return only the improved bullet points."
    )
    # prompt for experience enhancement
    openai_guider_1=(
        "You are a resume assistant."
        "Rewrite resume bullet points to be clear, precise and professional."
        "If the input points are short, increase the sentence length."
        "If the input points are long enough already, don't change anything except for grammatical errors"
        "Keep technical terms, metrics, and achievements intact."
        "Return only the improved bullet points."
    )
    for exp in input_data.get("experience", []):
        if exp.get("description"):
            bullets = "\n".join(exp["description"])
            response = client.chat.completions.create(
                model= "gpt-4o-mini",
                messages=[
                    {"role": "system", "content": openai_guider_1},
                    {"role": "user", "content": f"Polish these resume bullet points:\n{bullets}"}
                ]
            )
            improved = response.choices[0].message.content.strip().split("\n")
            exp["description"] = [b.strip("-• ") for b in improved if b.strip()]
    for proj in input_data.get("projects", []):
        if proj.get("description"):
            bullets = "\n".join(proj["description"])
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": openai_guider_1},
                    {"role": "user", "content": f"Polish these resume bullet points:\n{bullets}"}
                ]
            )
            improved = response.choices[0].message.content.strip().split("\n")
            proj["description"] = [b.strip("-• ") for b in improved if b.strip()]
    # prompt for skills enhancement
    if input_data.get("skills", {}).get("duties"):
        bullets = "\n".join(input_data["skills"]["duties"])
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": openai_guider_0},
                {"role": "user", "content": f"Polish these resume bullet points:\n{bullets}"}
            ]
        )
        improved = response.choices[0].message.content.strip().split("\n")
        input_data["skills"]["duties"] = [b.strip("-• ") for b in improved if b.strip()]

    return input_data

def add_education(doc, school, grad_date=None, degree=None, coursework=None):
    table=doc.add_table(rows=1,cols=2)
    table.autofit=False                
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.5)

    left=table.cell(0, 0).paragraphs[0] 
    right=table.cell(0, 1).paragraphs[0] 
    school=left.add_run(school)
    style_as(school, size=12, bold=True)
    if grad_date:
        right.alignment=align.RIGHT
        gdate=right.add_run(grad_date)
        style_as(gdate, size=12, italic=True)
    DC_para=doc.add_paragraph()
    deg=DC_para.add_run(degree)
    style_as(deg, size=12, bold=False)
    crse=DC_para.add_run(f"\nRelevant coursework: {coursework}")
    style_as(crse, size=12, italic=False)
def add_experience(doc, job, level, employer, date_range=None, description=None):
    table=doc.add_table(rows=1, cols=2)
    table.autofit=False                   
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.5)

    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    job1=left.add_run(job)
    style_as(job1, size=12, bold=True)
    job1.font.underline=True  
    left.add_run(" | ")
    level1=left.add_run(level)
    style_as(level1, size=12, bold=True)

    if date_range:
        right.alignment=align.RIGHT
        dRange=right.add_run(date_range)
        style_as(dRange, size=11, italic=True)
    Employer_para=doc.add_paragraph()
    Employer_para.paragraph_format.space_before = Pt(0)
    Employer_para.paragraph_format.space_after = Pt(0)
    emp_data=Employer_para.add_run(employer)
    style_as(emp_data, size=11, italic=True)
    if description:
        for d in description:
            job_des=doc.add_paragraph(d, style="List Bullet")
            job_des.paragraph_format.left_indent = Inches(0.5)
            job_des.paragraph_format.space_before = Pt(0)
            job_des.paragraph_format.space_after = Pt(0)
def add_project(doc, pr_cr_name, timeframe=None, description=None, framework=None, domain=None):
    table=doc.add_table(rows=1, cols=2)
    table.autofit=False                  
    table.columns[0].width = Inches(5.5)   
    table.columns[1].width = Inches(1.5)   
    left=table.cell(0, 0).paragraphs[0]
    right=table.cell(0, 1).paragraphs[0]
    PC_name=left.add_run(pr_cr_name)
    style_as(PC_name, size=12, bold=True)
    PC_name.font.underline=True
    if timeframe:
        right.alignment=align.RIGHT
        PC_time=right.add_run(timeframe)
        style_as(PC_time, size=11, italic=True)
    if description:
        for d in description:
            Discriotion_para=doc.add_paragraph(d, style="List Bullet")
            Discriotion_para.paragraph_format.left_indent = Inches(0.5)
            Discriotion_para.paragraph_format.space_before = Pt(0)
            Discriotion_para.paragraph_format.space_after = Pt(0)
    if framework or domain:
        fwork = ", ".join(framework) if framework else ""
        dmn = domain if domain else ""
        line = f"Framework: {fwork}"
        if dmn:
            line += f" | {dmn}"

        FDmn_para=doc.add_paragraph(line, style="List Bullet")
        FDmn_para.paragraph_format.left_indent = Inches(0.5)
        FDmn_para.paragraph_format.space_before = Pt(0)
        FDmn_para.paragraph_format.space_after = Pt(0)
def add_skills(doc, prog_lang, frm_lib, duties):
    para_skill=doc.add_paragraph()
    pl=para_skill.add_run("Programming Languages: ")
    style_as(pl, size=12, bold=True)  
    plnag=para_skill.add_run(" | ".join(prog_lang))
    style_as(plnag, size=10, bold=False)
    para_skill.add_run("\n")  
    fl=para_skill.add_run("Frameworks & Libraries: ")
    style_as(fl, size=12, bold=True)
    flib=para_skill.add_run(" | ".join(frm_lib))
    style_as(flib, size=10, bold=False)
    para_skill.add_run("\n")
    rr=para_skill.add_run("Roles & Responsibilities: ")
    style_as(rr, size=12, bold=True)
    rresp=para_skill.add_run(" | ".join(duties))
    style_as(rresp, size=10, bold=False)
    para_skill.paragraph_format.space_after = Pt(0)
    para_skill.paragraph_format.space_before = Pt(0)

def main():
    api=input("Enter a valid openAI api key:")
    if validate_api(api):
        print("API key is verified")
        client=OpenAI(api_key=api)
    else:
        sys.exit("Invalid API, exiting....")
    raw_data = build_resume() 
    enhanced_data = enhance_resume_data(raw_data,client)
    doc = Document()
    margins = doc.sections
    for margin in margins:
        margin.top_margin=Inches(0.5)
        margin.bottom_margin=Inches(0.5)
        margin.left_margin=Inches(0.5)
        margin.right_margin=Inches(0.5)
    add_head_info(doc, **enhanced_data["name_and_contact"])
    add_sec_head(doc, "Education")
    for edu in enhanced_data["education"]:
        add_education(doc, **edu)
    add_sec_head(doc, "Professional Experience")
    for exp in enhanced_data["experience"]:
        add_experience(doc, **exp)
    add_sec_head(doc, "Projects & Certifications")
    for proj in enhanced_data["projects"]:
        add_project(doc, **proj)

    add_sec_head(doc, "Technical Skills & Extra Curricular")
    add_skills(doc, **enhanced_data["skills"])
    doc.save("rsm_final.docx")


if __name__ == "__main__":
    main()

